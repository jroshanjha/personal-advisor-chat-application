from datetime import date
import streamlit as st
import requests
import pandas as pd
import os
from dotenv import load_dotenv
import google.generativeai as genai
import mysql.connector
import matplotlib.pyplot as plt
from PIL import Image
# for pdf reader or writer
from PyPDF2 import PdfReader,PdfMerger,PdfWriter
import docx
import json

load_dotenv()
## configure the api key:-
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-pro")

def get_gemini_response(text):
    response = model.generate_content(text)
    return response.text
m_chat = model.start_chat(history=[])
def get_gemini_response_store(question,k):
    if k=='m':
        response=m_chat.send_message(question,stream=True)
        return response  
def fetch_drug_info(drug_name):
    url = f"https://api.fda.gov/drug/label.json?search=openfda.brand_name:{drug_name}"
    response = requests.get(url)
    if response.status_code == 200:
        return response.json()
    else:
        return {"error": "Drug not found"}
 
def fetch_healthcare_providers(location, specialization):
    api_key = "your_google_api_key"
    url = f"https://maps.googleapis.com/maps/api/place/textsearch/json?query={specialization}+in+{location}&key={api_key}"
    response = requests.get(url)
    if response.status_code == 200:
        return response.json()
    else:
        return {"error": "No providers found"}
    
def extract_text_from_docx(uploaded_file):
    """
    Extract text from a .docx file.
    """
    doc = docx.Document(uploaded_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# def extract_text_from_docx(uploaded_file):
#     """
#     Extract text from a .docx file.
#     """
#     doc = docx.Document(uploaded_file)
#     text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
#     return text

def analyze_content(content, file_type=None):
    """
    Perform dynamic analysis of the provided text or data.
    """
    st.write("### Content Analysis")
    if file_type in ["txt", "docx", "pdf"] or isinstance(content, str):
        # Basic word and character count
        word_count = len(content.split())
        char_count = len(content)
        st.write(f"**Word Count:** {word_count}")
        st.write(f"**Character Count:** {char_count}")

        # Optional NLP features (e.g., keyword extraction)
        if st.button("Perform Keyword Analysis"):
            keywords = set(content.split())  # Simple placeholder for keyword extraction
            st.write("**Keywords:**", ", ".join(keywords))

    elif file_type in ["csv", "xlsx"]:
        st.write("**Data Summary**")
        st.write(content.describe())

    elif file_type == "json":
        st.write("**JSON Structure**")
        st.json(content)

    else:
        st.warning("Unsupported content type for analysis.")

def process_file(uploaded_file):
    """
    Dynamically processes the uploaded file and extracts content.
    """
    file_extension = uploaded_file.name.split(".")[-1].lower()
    extracted_content = None

    if file_extension == "pdf":
        st.write("### PDF Analysis")
        reader = PdfReader(uploaded_file)
        extracted_content = "\n".join(page.extract_text() for page in reader.pages)
        st.text_area("Extracted Text from PDF", extracted_content, height=200)

    elif file_extension in ["csv", "xlsx"]:
        st.write("### Spreadsheet Analysis")
        try:
            if file_extension == "csv":
                df = pd.read_csv(uploaded_file)
            elif file_extension == "xlsx":
                df = pd.read_excel(uploaded_file)
            st.write("Data Preview:")
            st.dataframe(df)
            extracted_content = df
        except Exception as e:
            st.error(f"Error reading the spreadsheet: {e}")

    elif file_extension == "txt":
        st.write("### Text File Analysis")
        extracted_content = uploaded_file.read().decode("utf-8")
        st.text_area("Text Content", extracted_content, height=200)

    elif file_extension == "docx":
        st.write("### Word Document Analysis (.docx)")
        extracted_content = extract_text_from_docx(uploaded_file)
        st.text_area("Extracted Text from DOCX", extracted_content, height=200)

    elif file_extension in ["jpg", "jpeg", "png", "gif", "bmp"]:
        st.write("### Image Analysis")
        try:
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded Image", use_column_width=True)
            extracted_content = "Image uploaded successfully."
        except Exception as e:
            st.error(f"Error processing the image: {e}")

    elif file_extension == "json":
        st.write("### JSON File Analysis")
        try:
            extracted_content = json.load(uploaded_file)
            st.json(extracted_content)
        except Exception as e:
            st.error(f"Error parsing JSON: {e}")

    else:
        st.write("### Unsupported File Type")
        extracted_content = uploaded_file.read()
        st.code(extracted_content, language="text")

    return extracted_content, file_extension

if 'mchat_history' not in st.session_state:
    st.session_state['mchat_history'] = []
    
def medical_method():
    # Title
    #st.set_page_config("AI-Powered Medical Diagnostic Tool 🩺")
    st.title("Welcome to Medical Diagnostic Tool 🩺")
    # Sidebar Menu
    menu = st.sidebar.radio("Menu", ["Symptom Diagnosis","File Summary","Drug Information", "Healthcare Finder"])
    prompt = "Tell me about Medical Diagnostic Tool"
    if menu == "Symptom Diagnosis":
        st.header("Symptom Diagnosis")
        symptoms = st.text_area("Enter your symptoms (e.g., fever, headache, fatigue)")
        
        if st.button("Get Diagnosis"):
            prompt = f"Given the symptoms: {symptoms}, provide a list of possible diagnoses and recommended actions."
            st.write("----------------------------------------------------------------")
            st.write(f"**Possible Diagnoses and Recommendations: for {symptoms}**")
            st.write("----------------------------------------------------------------")
            #response = get_gemini_response(prompt)
            response = get_gemini_response_store(prompt,'m')
            st.session_state['mchat_history'].append(("You", symptoms))
            st.subheader("The Response is")
            for chunk in response:
                st.write(chunk.text)
                st.session_state['mchat_history'].append(("Bot", chunk.text))
                
        st.subheader("The Chat History is")
        for role, text in st.session_state['mchat_history']:
            st.write(f"{role}: {text}")

    elif menu =='File Summary':
        # Streamlit App
        st.title("Dynamic File and Text Analysis")
        # Upload file
        uploaded_file = st.file_uploader("Upload any file", type=None)
        # User-provided text input
        #user_input = st.text_area("Enter additional text for analysis or comments", height=150)
        # Process uploaded file
        if uploaded_file:
            st.write("----------------------------------------------------------------")
            st.write(f"**Uploaded File:** {uploaded_file.name}")
            st.write("----------------------------------------------------------------")
            file_content, file_type = process_file(uploaded_file)

            # Perform analysis on file content
            if file_content:
                analyze_content(file_content, file_type)

        # Perform dynamic analysis on user input
        # if user_input.strip():
        #     st.write("### User Input Analysis")
        #     analyze_content(user_input, file_type="txt")
            
        # else:
        #     st.write("### No File Uploaded")
        
    elif menu == "Drug Information":
        st.header("Drug Information")
        drug_name = st.text_input("Enter the drug name")
        if st.button("Search Drug Info"):
            if drug_name:  # Ensure a drug name is entered
                data = fetch_drug_info(drug_name)
                prompt = f"Give me depth idea about Grug Information:- {drug_name}"
                # if "error" not in data:  # Check for errors in the API response
                #     st.write("**Drug Information:**")
                #     prompt = st.json(data)  # Display the drug data in JSON format
                # else:
                #     st.error(data.get("error", "Unable to fetch drug information."))
                st.write("----------------------------------------------------------------")
                st.write(f"**Drug Information: about {drug_name}**")
                st.write("----------------------------------------------------------------")
                # prompt = st.json(data)
                #st.write("Fetching drug information... (Replace this with an API call to OpenFDA)")
                # Example: Use OpenFDA API for real-time drug data
            else:
                st.warning("Please enter a drug name.")
                
            response = get_gemini_response(prompt)
            st.subheader('The response Output:-')
            st.write(response)        
        
    elif menu == "Healthcare Finder":
        st.header("Find Healthcare Providers")
        location = st.text_input("Enter your location")
        specialization = st.text_input("Enter specialization (e.g., cardiology, dermatology)")
        
        if st.button("Search Providers"):
            providers = fetch_healthcare_providers(location, specialization)
            prompt = f"Search My Health Care Finder for Locations:-{location} and specialication:- {specialization}"
            st.write("----------------------------------------------------------------")
            st.write("** Health Care Information:**")
            st.write("----------------------------------------------------------------")
            # for provider in providers.get("results", []):
            #     st.write(f"Name: {provider['name']}")
            #     st.write(f"Address: {provider['formatted_address']}")
            #     st.write("---")
        
            #st.write("Fetching healthcare providers... (Integrate with Zocdoc or similar API)")
            # Placeholder: Implement Zocdoc or similar API integration
        
            response = get_gemini_response(prompt)
            st.subheader('The response Output:-')
            st.write(response)
# Main
if __name__ == "__main__":
    medical_method()